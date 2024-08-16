import os
import json
import pandas as pd
from openai import AzureOpenAI
from typing import List, Dict
from docx import Document
from cosmos_db_service import cosmos_db_service
from dotenv import load_dotenv
from io import BytesIO

# Load environment variables from the .env file
load_dotenv()

api_key = os.getenv("AZURE_OPENAI_API_KEY")
api_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
local_resume_folder = os.getenv("LOCAL_RESUME_FOLDER")

client = AzureOpenAI(
    api_key=api_key,
    # per comment in staffing requirements extractor, when the 2024-08-06 API is available, we'll switch to that to use structured output
    api_version="2024-02-15-preview",
    azure_endpoint=api_endpoint
)

cosmos_db_service = cosmos_db_service()
cosmos_db_service.initialize()

def generate_resume_content(employee_data, staffing_data):
    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            {
                "role": "system",
                'content': """You are a helpful AI assistant that specializes in generating resumes for employees from a database of employee records.
                You will be provided with a JSON object that contains the employee's certifications, education, skills, and work history.
                The JSON object will have the following keys: certifications, education, skills, and employment_history.
                Each value in the JSON object will be a tab-separated string that contains the data for the respective category.
                Here is an example:
                {
                    "certifications": "Project Management Professional, 2011 - Present",
                    "education": "MPA, School of Public and Environmental Affairs, Indiana University, 2010",
                    "skills": "Data Mining",
                    "employment_history": "Acme Corporation, Senior Consultant, 2010 - Present",
                    "security_clearances": "Secret"
                }
                You will also be provided a JSON that has information about a position's requirements for education, certifications, experience, and security clearances.
                Here is an example:
                {
                    "required_role": "Program Manager (Key Personnel)",
                    "education": {
                        "Degree": "M.S.",
                        "Field": [
                            "Meteorology",
                            "Oceanography",
                            "Physical Science",
                            "Mathematics",
                            "Engineering"
                        ]
                    },
                    "certifications": [],
                    "experience": [
                        {"requirement": "Must have current knowledge and demonstrated experience with Navy METOC operations, operational systems, applications and requirements"},
                        {"requirement": "Must have at a minimum 5 years of demonstrated experience managing scientific, engineering, computing, and technical personnel working on Numerical Weather Prediction (NWP) and METOC applications development projects"}
                    ],
                    "security_clearance": []
                }
                Your task is to generate a JSON object that contains a resume for each employee that satisfies the requirements for the given positions. This should be a list of JSON objects where each item in the list represents a candidate resume. YOu s
                The JSON should have the following format:
                {
                    "name": The name of the employee. Do not include the employee id in the name.,
                    "employee_id": The employee's id.,
                    "professional_summary": This should be a string that summarizes the employee's professional experience. This summary should emphasize the employee's qualifications for the position.,
                    "key_competencies": This should be a list of strings from the employee's skills. Include only the skills that are relevant to the position.,
                    "education": This should be a list of strings in the format "Degree earned, Field of study, School Attended, Date of graduation",
                    "certifications": This should be a list of strings from the employee's certifications,
                    "relevant_experience": This should be a list of strings from the employment_history's responsibilitiesAndAchievements that match the experience requirements. Do not include employment history that is not relevant to the position,
                    "employment_history": This should be a list of strings in the format "Company, Title, Start Date - End Date",
                    "years_of_experience": The total number of years the employee has worked in the field. This should be calculated from the employment_history's startDate and endDate,
                    "years_of_relevant_experience": The total number of years the employee has worked in a relevant position. This should be calculated from the employment_history's startDate and endDate,
                    "security_clearances": This should be a list of strings from the security_clearances,
                }
                Do not include information you do not know. Do not include information about the employee that is not provided in the employee data.
                Your response should be in JSON format and include only the JSON."""
            },
            {
                "role": "user",
                "content": f"""Generate a list of candidate resumes using the following employee data: {employee_data}\n\n and find roles from this data: {staffing_data}"""
            }
        ],
        max_tokens=4000,
        seed=42
    )
    
    retry_count = 0
    json_data = {}
    cleaned_json = {}

    while (retry_count < 3):

        # if we are retrying, send back in the json_data for the LLM to retry the format
        if (retry_count > 0):
            response = client.chat.completions.create(
                model=deployment_name,
                messages=[
                    {
                        "role": "user",
                        "content": f"""The JSON you returned could not be sent into json.loads. Please take the following data and fix it: {cleaned_json}"""
                    }
                ],
                max_tokens=4000,
                seed=42
            )
    
        message_content = response.choices[0].message.content.strip()
            
        # Find the JSON content in the response
        try:
            message_content = str(message_content).split("```")[1]

            if message_content.startswith('json'):
                # Remove the first occurrence of 'json' from the response text
                message_content = message_content[4:]

            cleaned_json = message_content.replace('\n', '').replace('\\n', '').replace('\\', '').strip('"')

            json_data = json.loads(cleaned_json)

            return json_data
        except (json.JSONDecodeError, ValueError, IndexError) as e:
            print("Error: The response content is not valid JSON")
            retry_count = retry_count + 1

    print("Retry count exceeded!")

def create_resume(json_matches):
    print(json.dumps(json_matches, indent=4))

    with open('moqdata/ResumeTemplate.docx', 'rb') as file:
        byte_stream = file.read()

    byte_stream_io = BytesIO(byte_stream)

    document = Document(byte_stream_io)

    name = json_matches['name']
    education = json_matches["education"]
    key_competencies = json_matches["key_competencies"]
    certifications = json_matches["certifications"]
    professional_summary = json_matches["professional_summary"]
    relevant_experience = json_matches['relevant_experience']
    employment_history = json_matches['employment_history']
    security_clearances = json_matches['security_clearances']
     
    for paragraph in document.paragraphs:
        if 'Name' in paragraph.text:
            paragraph.text = paragraph.text.replace('Name', name)
            paragraph.text += "\n\n" + professional_summary
    
    for table in document.tables:
        for row in table.rows:
            # Check the leftmost cell (first cell) for the search text
            left_cell = row.cells[0]
            for paragraph in left_cell.paragraphs:
                if 'Key Competencies' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        key_competencies_str = "\n\n".join(key_competencies)
                        right_paragraph.text = key_competencies_str
                        break

                if 'Education' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        education_str = "\n\n".join(education)
                        right_paragraph.text = education_str
                        break

                if 'Training & Certifications' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        certifications_str = "\n\n".join(certifications)
                        right_paragraph.text = certifications_str
                        break
              
                if 'Security Clearances' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        security_clearances_str = "\n\n".join(security_clearances)
                        right_paragraph.text = security_clearances_str
                        break

                if 'Experience' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        relevant_experience_str = "\n\n".join(relevant_experience)
                        right_paragraph.text = relevant_experience_str
                        break

                if 'History' in paragraph.text:
                    # If found, set the text in the rightmost cell
                    right_cell = row.cells[-1]
                    for right_paragraph in right_cell.paragraphs:
                        employment_history_str = "\n\n".join(employment_history)
                        right_paragraph.text = employment_history_str
                        break

    resume_name_path =local_resume_folder + "\\" + name + "_Resume.docx"
    document.save(resume_name_path)
    print(f"Resume created: {resume_name_path}")

# get moq employee data, assumes a detailed search
# employee_data = None
# with open('moqdata/employeeData.json', 'r') as file:
#     employee_data = json.load(file)

csv_files = [file for file in os.listdir('moqdata') if file.endswith('.csv')]

dataframes = []
for file in csv_files:
    df = pd.read_csv(os.path.join('moqdata', file))
    dataframes.append(df)

merged_df = pd.concat(dataframes, ignore_index=True)
employee_data = merged_df.to_json(orient='records')

# get grouped rfp staffing data
grouped_staffing_data = cosmos_db_service.get_grouped_rfp_staffing_extract()

# send employee data & rfp staffing data to OpenAI for processing
json_matches = generate_resume_content(json.dumps(employee_data), json.dumps(grouped_staffing_data))

if (json_matches):
    for match in json_matches:
        create_resume(match)